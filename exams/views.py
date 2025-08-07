from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import IntegrityError
from django.contrib.auth.decorators import user_passes_test
from django.core.paginator import Paginator
from django.db.models import Q, Count
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.core.mail import send_mail
from django.conf import settings
from .models import Course, Unit, Student, ExamRecord, Campus, CampusPassword, School
from .forms import ExamRecordForm, CourseForm, UnitForm, StudentForm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import json
import io


def is_superuser(user):
    return user.is_superuser


def campus_select(request):
    campuses = Campus.objects.all()
    if request.method == 'POST':
        campus_id = request.POST.get('campus_id')
        
        if campus_id:
            campus = get_object_or_404(Campus, id=campus_id)
            password = request.POST.get(f'password_{campus_id}', '')
            
            # Check if campus has a password set
            try:
                campus_password = CampusPassword.objects.get(campus=campus)
                
                if not campus_password.password:
                    messages.error(request, f'Access denied. {campus.name} requires a password set by administrator.')
                    return render(request, 'campus_select.html', {'campuses': campuses})
                
                if campus_password.password != password:
                    messages.error(request, 'Invalid password for this campus.')
                    return render(request, 'campus_select.html', {'campuses': campuses})
                    
            except CampusPassword.DoesNotExist:
                # No password set for this campus - deny access
                messages.error(request, f'Access denied. {campus.name} requires a password set by administrator.')
                return render(request, 'campus_select.html', {'campuses': campuses})
            
            request.session['campus_id'] = int(campus_id)
            return redirect('exams:home')
    
    return render(request, 'campus_select.html', {'campuses': campuses})


@login_required
@user_passes_test(is_superuser)
def manage_campus_passwords(request):
    # Double-check that only superusers can access this
    if not request.user.is_superuser:
        messages.error(request, 'Access denied. Only administrators can manage campus passwords.')
        return redirect('exams:home')
    
    if request.method == 'POST':
        campus_id = request.POST.get('campus_id')
        
        if campus_id:
            campus = get_object_or_404(Campus, id=campus_id)
            password = request.POST.get(f'password_{campus_id}', '').strip()
            
            # Ensure password is not empty
            if not password:
                messages.error(request, f'Password is required for {campus.name}. Cannot leave empty.')
                return redirect('exams:manage_campus_passwords')
            
            campus_password, created = CampusPassword.objects.get_or_create(campus=campus)
            campus_password.password = password
            campus_password.save()
            messages.success(request, f'Password updated for {campus.name}')
        
        return redirect('exams:manage_campus_passwords')
    
    campuses = Campus.objects.all()
    campus_passwords = {}
    for campus in campuses:
        try:
            campus_passwords[campus.id] = CampusPassword.objects.get(campus=campus).password
        except CampusPassword.DoesNotExist:
            campus_passwords[campus.id] = ''
    
    context = {
        'campuses': campuses,
        'campus_passwords': campus_passwords
    }
    return render(request, 'exams/manage_campus_passwords.html', context)


def get_current_campus(request):
    campus_id = request.session.get('campus_id')
    if campus_id:
        return Campus.objects.get(id=campus_id)
    return None 


def home(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Get counts for dashboard - respect campus selection for both superusers and regular users
    if current_campus:
        total_students = Student.objects.filter(course__school__campus=current_campus).count()
        total_units = Unit.objects.filter(course__school__campus=current_campus).count()
        total_records = ExamRecord.objects.filter(student__course__school__campus=current_campus).count()
        total_courses = Course.objects.filter(school__campus=current_campus).count()
        recent_records = ExamRecord.objects.select_related('student', 'unit').filter(student__course__school__campus=current_campus).order_by('-id')[:5]
    else:
        # If no campus is selected and user is superuser, show all records
        if request.user.is_superuser:
            total_students = Student.objects.count()
            total_units = Unit.objects.count()
            total_records = ExamRecord.objects.count()
            total_courses = Course.objects.count()
            recent_records = ExamRecord.objects.select_related('student', 'unit').order_by('-id')[:5]
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    context = {
        'total_students': total_students,
        'total_units': total_units,
        'total_records': total_records,
        'total_courses': total_courses,
        'recent_records': recent_records,
        'current_campus': current_campus,
    }
    return render(request, 'exams/home.html', context)


def enter_marks(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    if current_campus:
        students = Student.objects.filter(course__school__campus=current_campus)
        courses = Course.objects.filter(school__campus=current_campus)
        units = Unit.objects.filter(course__school__campus=current_campus)
        schools = School.objects.filter(campus=current_campus)
    else:
        if request.user.is_superuser:
            students = Student.objects.all()
            courses = Course.objects.all()
            units = Unit.objects.all()
            schools = School.objects.all()
        else:
            return redirect('exams:campus_select')
    years = [2025, 2024, 2023, 2022, 2021]
    terms = ['Term 1', 'Term 2', 'Term 3', 'Term 4']
    selected_student = None
    selected_year = None
    selected_term = None
    unit_marks = []
    available_units = []
    message = ''
    if request.method == 'POST' and 'select_student' in request.POST:
        student_id = request.POST.get('student_id')
        year = request.POST.get('year')
        term = request.POST.get('term')
        selected_student = get_object_or_404(Student, id=student_id)
        selected_year = year
        selected_term = term
        all_units = list(Unit.objects.filter(course=selected_student.course))
        entered_units = []
        for unit in all_units:
            rec = ExamRecord.objects.filter(student=selected_student, unit=unit, year=year, term=term).first()
            if rec:
                unit_marks.append({
                    'unit': unit,
                    'cat1': rec.cat1_score,
                    'cat2': rec.cat2_score,
                    'endterm': rec.end_term_score,
                    'has_record': True,
                })
                entered_units.append(unit.id)
        available_units = [u for u in all_units if u.id not in entered_units]
    elif request.method == 'POST' and 'save_marks' in request.POST:
        student_id = request.POST.get('student_id')
        year = request.POST.get('year')
        term = request.POST.get('term')
        selected_student = get_object_or_404(Student, id=student_id)
        selected_year = year
        selected_term = term
        school = request.POST.get('school')
        level = request.POST.get('level')
        i = 1
        entered_units = []
        while True:
            unit_id = request.POST.get(f'unit_id_{i}')
            unit_name = request.POST.get(f'unit_name_{i}', '').strip()
            if not unit_id:
                break
            cat1 = request.POST.get(f'cat1_{i}')
            cat2 = request.POST.get(f'cat2_{i}')
            endterm = request.POST.get(f'endterm_{i}')
            cat1_val = safe_decimal(cat1)
            cat2_val = safe_decimal(cat2)
            endterm_val = safe_decimal(endterm)
            if unit_id == 'other' and unit_name:
                unit, _ = Unit.objects.get_or_create(
                    name=unit_name,
                    course=selected_student.course
                )
            else:
                unit = Unit.objects.get(id=unit_id)
            record, _ = ExamRecord.objects.get_or_create(
                student=selected_student,
                unit=unit,
                year=year,
                term=term,
                defaults={
                    'cat1_score': cat1_val,
                    'cat2_score': cat2_val,
                    'end_term_score': endterm_val,
                }
            )
            record.cat1_score = cat1_val
            record.cat2_score = cat2_val
            record.end_term_score = endterm_val
            record.save()
            entered_units.append(unit.id)
            i += 1
        message = 'Marks saved successfully!'
        unit_marks = []
        all_units = list(Unit.objects.filter(course=selected_student.course))
        for unit in all_units:
            rec = ExamRecord.objects.filter(student=selected_student, unit=unit, year=year, term=term).first()
            if rec:
                unit_marks.append({
                    'unit': unit,
                    'cat1': rec.cat1_score,
                    'cat2': rec.cat2_score,
                    'endterm': rec.end_term_score,
                    'has_record': True,
                })
        available_units = [u for u in all_units if u.id not in entered_units]
    context = {
        'students': students,
        'courses': courses,
        'units': units,
        'schools': schools,
        'years': years,
        'terms': terms,
        'selected_student': selected_student,
        'selected_year': selected_year,
        'selected_term': selected_term,
        'unit_marks': unit_marks,
        'available_units': available_units,
        'message': message,
        'current_campus': current_campus,
    }
    return render(request, 'exams/enter_marks.html', context)


def get_existing_marks(request):
    """AJAX endpoint to get existing marks for students"""
    if request.method == 'GET':
        course_id = request.GET.get('course_id')
        unit_id = request.GET.get('unit_id')
        term = request.GET.get('term')
        year = request.GET.get('year')
        
        if not all([course_id, unit_id, term, year]):
            return JsonResponse({'error': 'Missing required parameters'}, status=400)
        
        try:
            course = Course.objects.get(id=course_id)
            unit = Unit.objects.get(id=unit_id)
            
            # Get students for this course
            students = Student.objects.filter(course=course)
            
            # Get existing marks for these students for the specified unit, term, and year
            existing_marks = {}
            for student in students:
                record = ExamRecord.objects.filter(
                    student=student,
                    unit=unit,
                    term=term,
                    year=year
                ).first()
                
                if record:
                    existing_marks[student.id] = {
                        'cat1_score': float(record.cat1_score) if record.cat1_score else '',
                        'cat2_score': float(record.cat2_score) if record.cat2_score else '',
                        'end_term_score': float(record.end_term_score) if record.end_term_score else ''
                    }
                else:
                    existing_marks[student.id] = {
                        'cat1_score': '',
                        'cat2_score': '',
                        'end_term_score': ''
                    }
            
            return JsonResponse({'existing_marks': existing_marks})
            
        except (Course.DoesNotExist, Unit.DoesNotExist):
            return JsonResponse({'error': 'Invalid course or unit'}, status=400)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)


def view_records(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # For both superusers and regular users, filter by campus if one is selected
    if current_campus:
        records = ExamRecord.objects.select_related('student', 'unit', 'unit__course').filter(student__course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all records
        if request.user.is_superuser:
            records = ExamRecord.objects.select_related('student', 'unit', 'unit__course').all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    # Filtering
    student_filter = request.GET.get('student')
    course_filter = request.GET.get('course')
    unit_filter = request.GET.get('unit')
    term_filter = request.GET.get('term')
    year_filter = request.GET.get('year')
    
    if student_filter:
        records = records.filter(
            Q(student__name__icontains=student_filter) | 
            Q(student__registration_number__icontains=student_filter)
        )
    if course_filter:
        records = records.filter(unit__course__name__icontains=course_filter)
    if unit_filter:
        records = records.filter(unit__name__icontains=unit_filter)
    if term_filter:
        records = records.filter(term__iexact=term_filter)
    if year_filter:
        records = records.filter(year=year_filter)
    
    # Pagination
    paginator = Paginator(records, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # For filter dropdowns - respect campus selection
    if current_campus:
        years = ExamRecord.objects.filter(student__course__school__campus=current_campus).values_list('year', flat=True).distinct().order_by('year')
        terms = ExamRecord.objects.filter(student__course__school__campus=current_campus).values_list('term', flat=True).distinct().order_by('term')
        courses = Course.objects.filter(school__campus=current_campus)
        units = Unit.objects.filter(course__school__campus=current_campus)
        students = Student.objects.filter(course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            years = ExamRecord.objects.values_list('year', flat=True).distinct().order_by('year')
            terms = ExamRecord.objects.values_list('term', flat=True).distinct().order_by('term')
            courses = Course.objects.all()
            units = Unit.objects.all()
            students = Student.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    context = {
        'page_obj': page_obj,
        'courses': courses,
        'units': units,
        'students': students,
        'years': years,
        'terms': terms,
        'filters': {
            'course': course_filter,
            'unit': unit_filter,
            'student': student_filter,
            'year': year_filter,
            'term': term_filter,
            'sort': 'student__name', # Default sort
        },
        'current_campus': current_campus,
    }
    return render(request, 'exams/view_records.html', context)


def update_record(request, record_id):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    record = get_object_or_404(ExamRecord, id=record_id)
    
    if request.method == 'POST':
        form = ExamRecordForm(request.POST, instance=record)
        if form.is_valid():
            form.save()
            messages.success(request, f'Exam record updated for {record.student.name} - {record.unit.name}')
            return redirect('exams:view_records')
    else:
        form = ExamRecordForm(instance=record)
    
    context = {
        'form': form,
        'record': record,
        'current_campus': current_campus,
    }
    return render(request, 'exams/update_record.html', context)


def manage_courses(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Get all schools with their campuses
    schools = School.objects.select_related('campus').all().order_by('name')
    
    # Get comprehensive course data with all relationships
    courses = Course.objects.select_related(
        'school',
        'school__campus'
    ).prefetch_related(
        'units',
        'students',
        'units__exam_records'
    ).annotate(
        unit_count=Count('units'),
        student_count=Count('students')
    ).order_by('-created_at')

    # Detailed debug logging
    print("\n=== COURSE ANALYSIS DEBUG ===")
    print(f"Total courses found: {courses.count()}")
    for i, course in enumerate(courses, 1):
        print(f"\nCourse #{i}: {course.name}")
        print(f"School: {course.school.name if course.school else 'None'}")
        print(f"Campus: {course.school.campus.name if course.school and course.school.campus else 'None'}")
        print(f"Units: {course.unit_count}")
        print(f"Students: {course.student_count}")
        if course.units.exists():
            print("Unit Details:")
            for unit in course.units.all():
                print(f"  - {unit.name}: {unit.exam_records.count()} records")

    # Apply campus filter if a campus is selected
    if current_campus:
        courses = courses.filter(school__campus=current_campus)
        schools = schools.filter(campus=current_campus)
        print(f"\nFiltered to {current_campus.name} campus: {courses.count()} courses")
    
    # Handle course creation
    if request.method == 'POST' and 'create_course' in request.POST:
        form = CourseForm(request.POST)
        if form.is_valid():
            course = form.save(commit=False)
            school_id = request.POST.get('school')
            
            if not school_id:
                messages.error(request, 'Please select a school')
                return redirect('exams:manage_courses')
                
            try:
                school = School.objects.get(id=school_id)
                
                # Check if course with same name already exists in this school
                existing_course = Course.objects.filter(
                    name__iexact=course.name,
                    school=school
                ).first()
                
                if existing_course:
                    messages.error(request, f'Course "{course.name}" already exists in {school.name}')
                    return redirect('exams:manage_courses')
                
                course.school = school  # Explicitly set school relationship
                course.save()
                
                # Debug output
                print(f"\nDEBUG: Created course - Name: {course.name}")
                print(f"School: {course.school.name if course.school else 'None'}")
                print(f"Campus: {course.school.campus.name if course.school and course.school.campus else 'None'}")
                print(f"Created at: {course.created_at}")
                
                messages.success(request, f'Course "{course.name}" created successfully for {school.name}!')
                return redirect('exams:manage_courses')
                
            except School.DoesNotExist:
                messages.error(request, 'Selected school does not exist')
                return redirect('exams:manage_courses')
            except IntegrityError:
                messages.error(request, f'Course "{course.name}" already exists in {school.name}')
                return redirect('exams:manage_courses')
    
    # Handle school creation
    elif request.method == 'POST' and 'create_school' in request.POST:
        school_name = request.POST.get('school_name', '').strip()
        if school_name:
            from django.db import transaction
            
            try:
                # First ensure we have a valid campus
                campus = current_campus or Campus.objects.first()
                if not campus:
                    messages.error(request, 'No campus available. Please create a campus first.')
                    return redirect('exams:manage_courses')

                with transaction.atomic():
                    # Check for existing school (case-insensitive) with same name AND campus
                    existing_school = School.objects.filter(
                        name__iexact=school_name,
                        campus=campus
                    ).first()
                    
                    if existing_school:
                        messages.error(request, f'School "{school_name}" already exists for {campus.name} campus')
                    else:
                        # Create new school
                        school = School.objects.create(
                            name=school_name,
                            campus=campus
                        )
                        messages.success(request, f'School "{school.name}" created successfully for {campus.name} campus!')
                
                return redirect('exams:manage_courses')
                
            except IntegrityError as e:
                # Check if school already exists for this campus
                existing = School.objects.filter(name__iexact=school_name, campus=campus).exists()
                if existing:
                    messages.error(request, f'School "{school_name}" already exists for {campus.name} campus')
                else:
                    messages.error(request, 'Database error occurred. Please try again with a different name.')
                return redirect('exams:manage_courses')
            except Exception as e:
                import traceback
                print(f"Error creating school: {str(e)}\n{traceback.format_exc()}")
                messages.error(request, 'An unexpected error occurred. Please try again.')
                return redirect('exams:manage_courses')
        else:
            messages.error(request, 'School name is required')
            return redirect('exams:manage_courses')
    
    # Handle course deletion
    elif request.method == 'POST' and 'delete_course' in request.POST:
        course_id = request.POST.get('course_id')
        try:
            course = Course.objects.get(id=course_id)
            course_name = course.name
            course.delete()
            messages.success(request, f'Course "{course_name}" deleted successfully!')
        except Course.DoesNotExist:
            messages.error(request, 'Course not found')
        return redirect('exams:manage_courses')
            
    else:
        form = CourseForm()
    
    # Debug output for template context
    print("\n=== TEMPLATE CONTEXT DEBUG ===")
    print(f"Passing {courses.count()} courses to template")
    for i, course in enumerate(courses, 1):
        print(f"Course {i}: {course.name} (ID: {course.id})")
        print(f"  School: {course.school.name if course.school else 'None'}")
        print(f"  Units: {course.units.count()}")
        print(f"  Students: {course.students.count()}")
    
    context = {
        'form': form,
        'courses': courses,
        'schools': schools,
        'current_campus': current_campus,
    }
    return render(request, 'exams/manage_courses.html', context)


def manage_units(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        units = Unit.objects.select_related('course').filter(course__school__campus=current_campus)
        courses = Course.objects.filter(school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            units = Unit.objects.select_related('course').all()
            courses = Course.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    if request.method == 'POST':
        form = UnitForm(request.POST)
        if form.is_valid():
            unit = form.save()
            messages.success(request, 'Unit created successfully!')
            return redirect('exams:manage_units')
    else:
        form = UnitForm()
    
    # Filter form choices based on campus
    if not request.user.is_superuser and current_campus:
        form.fields['course'].queryset = courses
    
    context = {
        'form': form,
        'units': units,
        'courses': courses,
        'current_campus': current_campus,
    }
    return render(request, 'exams/manage_units.html', context)


def manage_students(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        students = Student.objects.select_related('course').filter(course__school__campus=current_campus)
        courses = Course.objects.filter(school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            students = Student.objects.select_related('course').all()
            courses = Course.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    if request.method == 'POST':
        form = StudentForm(request.POST)
        if form.is_valid():
            student = form.save()
            messages.success(request, 'Student created successfully!')
            return redirect('exams:manage_students')
    else:
        form = StudentForm()
    
    # Filter form choices based on campus
    if not request.user.is_superuser and current_campus:
        form.fields['course'].queryset = courses
    
    context = {
        'form': form,
        'students': students,
        'courses': courses,
        'current_campus': current_campus,
    }
    return render(request, 'exams/manage_students.html', context)


def update_student(request, student_id):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    student = get_object_or_404(Student, id=student_id)
    
    if request.method == 'POST':
        form = StudentForm(request.POST, instance=student)
        if form.is_valid():
            form.save()
            messages.success(request, 'Student information updated successfully!')
            return redirect('exams:manage_students')
    else:
        form = StudentForm(instance=student)
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        courses = Course.objects.filter(school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            courses = Course.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    context = {
        'form': form,
        'student': student,
        'courses': courses,
        'current_campus': current_campus,
    }
    return render(request, 'exams/update_student.html', context)


def delete_record(request, record_id):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    record = get_object_or_404(ExamRecord, id=record_id)
    if request.method == 'POST':
        student_name = record.student.name
        unit_name = record.unit.name
        record.delete()
        messages.success(request, f'Exam record deleted for {student_name} - {unit_name}')
        return redirect('exams:view_records')
    
    context = {
        'record': record,
        'current_campus': current_campus,
    }
    return render(request, 'exams/delete_record.html', context)


def generate_report(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        courses = Course.objects.filter(school__campus=current_campus)
        students = Student.objects.filter(course__school__campus=current_campus)
        years = ExamRecord.objects.filter(student__course__school__campus=current_campus).values_list('year', flat=True).distinct().order_by('year')
        terms = ExamRecord.objects.filter(student__course__school__campus=current_campus).values_list('term', flat=True).distinct().order_by('term')
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            courses = Course.objects.all()
            students = Student.objects.all()
            years = ExamRecord.objects.values_list('year', flat=True).distinct().order_by('year')
            terms = ExamRecord.objects.values_list('term', flat=True).distinct().order_by('term')
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    pass_list = []
    selected_course = None
    selected_term = None
    selected_year = None
    transcript_preview = None
    pass_debug = []
    
    if request.method == 'POST' and 'student' in request.POST:
        student_id = request.POST.get('student')
        year = request.POST.get('year')
        term = request.POST.get('term')
        student = Student.objects.get(id=student_id)
        records = ExamRecord.objects.filter(student=student, year=year, term=term).select_related('unit', 'unit__course')
        
        if not records.exists():
            messages.error(request, 'No records found for this student, year, and term.')
        else:
            mean_total = 0
            results = []
            for rec in records:
                cat_avg = int(round((float(rec.cat1_score) + float(rec.cat2_score)) / 2))
                end_term = int(round(float(rec.end_term_score)))
                avg = int(round(cat_avg + end_term))
                if avg >= 75:
                    remark = 'Distinction'
                elif avg >= 60:
                    remark = 'Credit'
                elif avg >= 40:
                    remark = 'Pass'
                else:
                    remark = 'Fail'
                results.append({
                    'unit': rec.unit.name,
                    'cat1': rec.cat1_score,
                    'cat2': rec.cat2_score,
                    'cat': cat_avg,
                    'end_term': rec.end_term_score,
                    'average': avg,
                    'remark': remark,
                })
                mean_total += avg
            mean_score = mean_total // len(records) if records else 0
            transcript_preview = {
                'student': student,
                'year': year,
                'term': term,
                'results': results,
                'mean_score': mean_score,
            }
    
    context = {
        'courses': courses,
        'students': students,
        'years': years,
        'terms': terms,
        'transcript_preview': transcript_preview,
        'current_campus': current_campus,
        'dashboard_url': reverse('exams:campus_select'),
    }
    return render(request, 'exams/generate_report.html', context)


def download_report(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Handle POST request from generate_report form
    if request.method == 'POST':
        student_id = request.POST.get('student_id')
        year = request.POST.get('year')
        term = request.POST.get('term')
        
        if not all([student_id, year, term]):
            messages.error(request, 'Missing required parameters for report generation.')
            return redirect('exams:generate_report')
        
        try:
            student = Student.objects.get(id=student_id)
            records = ExamRecord.objects.filter(student=student, year=year, term=term).select_related('unit', 'unit__course')
            
            if not records.exists():
                messages.error(request, 'No records found for this student, year, and term.')
                return redirect('exams:generate_report')
            
            # Get student details
            student_name = student.name
            admission_number = student.registration_number
            course_name = student.course.name
            
        except Student.DoesNotExist:
            messages.error(request, 'Student not found.')
            return redirect('exams:generate_report')
    
    # Handle GET request (legacy filtering - for backward compatibility)
    else:
        if request.user.is_superuser:
            records = ExamRecord.objects.all()
        else:
            records = ExamRecord.objects.filter(student__course__school__campus=current_campus)
        
        # Filtering
        student_filter = request.GET.get('student')
        course_filter = request.GET.get('course')
        unit_filter = request.GET.get('unit')
        term_filter = request.GET.get('term')
        year_filter = request.GET.get('year')
        
        if student_filter:
            records = records.filter(student__name__icontains=student_filter)
        if course_filter:
            records = records.filter(unit__course__name__icontains=course_filter)
        if unit_filter:
            records = records.filter(unit__name__icontains=unit_filter)
        if term_filter:
            records = records.filter(term=term_filter)
        if year_filter:
            records = records.filter(year=year_filter)
        
        # Get details from filters or first record
        student_name = student_filter or ''
        admission_number = ''
        course_name = course_filter or ''
        year = year_filter or ''
        term = term_filter or ''
        school = ''
        level = ''
        
        if records.exists():
            first_record = records.first()
            if not student_name:
                student_name = first_record.student.name
            if not admission_number:
                admission_number = first_record.student.registration_number
            if not course_name:
                course_name = first_record.unit.course.name
            if not year:
                year = str(first_record.year)
            if not term:
                term = first_record.term
    
    # Create Word document
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    # HEADER IMAGE (full width, very top)
    try:
        doc.add_picture('report/head.jpg', width=Inches(7.0))
    except FileNotFoundError:
        doc.add_paragraph("Header image not found.")



    # TITLE
    p = doc.add_paragraph()
    run = p.add_run('STUDENTS   PROGRESS   REPORT')
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # STUDENT INFO (no table, just paragraphs, bold field names only)
    p = doc.add_paragraph()
    p.add_run("STUDENT'S NAME:").bold = True
    p.add_run(f" {student_name or '....................................................'}    ")
    p.add_run("ADM NO:").bold = True
    p.add_run(f" {admission_number or '................'}")
    p = doc.add_paragraph()
    p.add_run("COURSE:").bold = True
    p.add_run(f" {course_name or '..................................................'}    ")
    p.add_run("ACADEMIC YEAR:").bold = True
    p.add_run(f" {year or '................'}")
    p = doc.add_paragraph()
    p.add_run("TERM:").bold = True
    p.add_run(f" {term or '................'}")
    doc.add_paragraph()
    # RESULTS TABLE
    results_table = doc.add_table(rows=1, cols=5)
    results_table.style = 'Table Grid'
    hdr = results_table.rows[0].cells
    hdr[0].text = 'SUBJECT/UNIT'
    hdr[1].text = 'CAT'
    hdr[2].text = 'END TERM'
    hdr[3].text = 'AVERAGE%'
    hdr[4].text = 'REMARKS'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mean_total = 0
    for rec in records:
        row = results_table.add_row().cells
        row[0].text = rec.unit.name.upper()
        cat_avg = int(round((float(rec.cat1_score) + float(rec.cat2_score)) / 2))
        row[1].text = str(cat_avg)
        end_term = int(round(float(rec.end_term_score)))
        row[2].text = str(end_term)
        avg = int(round(cat_avg + end_term))
        row[3].text = str(avg)
        if avg >= 75:
            remark = 'Distinction'
        elif avg >= 60:
            remark = 'Credit'
        elif avg >= 40:
            remark = 'Pass'
        else:
            remark = 'Fail'
        row[4].text = remark
        mean_total += avg
        for i in range(1,5):
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mean_score = int(round(mean_total / records.count())) if records.count() else 0
    if mean_score >= 75:
        grade = 'A..(DISTINCTION)'
    elif mean_score >= 60:
        grade = 'B..(CREDIT)'
    elif mean_score >= 40:
        grade = 'C..(PASS )'
    else:
        grade = 'D..(FAIL)'
    doc.add_paragraph()
    # MEAN SCORE & GRADE (single paragraph, spaced)
    mean_grade_para = doc.add_paragraph()
    mean_grade_para.add_run(f'MEAN SCORE: {mean_score}').bold = True
    mean_grade_para.add_run(' ' * 15)
    mean_grade_para.add_run(f'GRADES: {grade}').bold = True
    doc.add_paragraph()
    # GRADING SYSTEM BOX (narrower, only grade letter bold)
    grading_table = doc.add_table(rows=5, cols=1)
    grading_table.style = 'Table Grid'
    grading_table.autofit = False
    grading_table.columns[0].width = Inches(1.2)
    grading_table.cell(0,0).text = ''
    grading_table.cell(0,0).paragraphs[0].add_run('Grading system:').bold = True
    # Row 1
    p = grading_table.cell(1,0).paragraphs[0]
    p.add_run('75-100 – ').bold = False
    p.add_run('A').bold = True
    p.add_run(' (Distinction)').bold = False
    # Row 2
    p = grading_table.cell(2,0).paragraphs[0]
    p.add_run('60 – 75 – ').bold = False
    p.add_run('B').bold = True
    p.add_run(' (Credit)').bold = False
    # Row 3
    p = grading_table.cell(3,0).paragraphs[0]
    p.add_run('40 – 59 – ').bold = False
    p.add_run('C').bold = True
    p.add_run(' (Pass)').bold = False
    # Row 4
    p = grading_table.cell(4,0).paragraphs[0]
    p.add_run('0 – 39 – ').bold = False
    p.add_run('D').bold = True
    p.add_run(' (Fail)').bold = False
    doc.add_paragraph()
    # SIGNATURES (not in a table, just paragraphs)
    sig_line = doc.add_paragraph()
    sig_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_line.add_run('Signed_________________').italic = True
    sig_line.add_run(' ' * 10)
    sig_line.add_run('Signed_________________').italic = True
    sig_titles = doc.add_paragraph()
    sig_titles.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_titles.add_run('EXAMINATION').bold = True
    sig_titles.add_run(' ' * 20)
    sig_titles.add_run('PRINCIPAL').bold = True
    for run in sig_titles.runs:
        run.italic = True
    doc.add_paragraph()
    # FOOTER IMAGE (full width, very bottom)
    doc.add_picture('report/foot.png', width=Inches(7.0))
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Progress_Report_{student_name.replace(' ', '_') if student_name else 'all'}_{year or 'all'}_T{term or 'all'}.docx"
    response['Content-Disposition'] = f'attachment; filename=\"{filename}\"'
    return response


def pass_list(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        records = ExamRecord.objects.filter(student__course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            records = ExamRecord.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    # Calculate averages and filter passing students (total average >= 40)
    students_passed = {}
    for record in records:
        # Skip records with incomplete scores
        if not all([record.cat1_score, record.cat2_score, record.end_term_score]):
            continue
            
        # Calculate total average for this record
        total_avg = float(record.cat_average + record.end_term_score)
        
        # Only include passing records (total_avg >= 40)
        if total_avg >= 40:
            student = record.student
            if student.id not in students_passed:
                students_passed[student.id] = {
                    'student': student,
                    'records': [],
                    'average': 0,
                    'count': 0,
                    'total': 0
                }
            students_passed[student.id]['records'].append(record)
            students_passed[student.id]['count'] += 1
            students_passed[student.id]['total'] += total_avg
    
    # Calculate overall average for each student
    for student_data in students_passed.values():
        student_data['average'] = student_data['total'] / student_data['count'] if student_data['count'] > 0 else 0
    
    # Sort by average
    sorted_students = sorted(students_passed.values(), key=lambda x: x['average'], reverse=True)
    
    context = {
        'students_passed': sorted_students,
        'current_campus': current_campus,
        'dashboard_url': reverse('exams:campus_select'),
    }
    return render(request, 'exams/pass_list.html', context)


def download_pass_list(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        records = ExamRecord.objects.filter(student__course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            records = ExamRecord.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    # Calculate averages and filter passing students (total average >= 40)
    students_passed = {}
    for record in records:
        # Skip records with incomplete scores
        if not all([record.cat1_score, record.cat2_score, record.end_term_score]):
            continue
            
        # Calculate total average for this record
        total_avg = float(record.cat_average + record.end_term_score)
        
        # Only include passing records (total_avg >= 40)
        if total_avg >= 40:
            student = record.student
            if student.id not in students_passed:
                students_passed[student.id] = {
                    'student': student,
                    'records': [],
                    'average': 0,
                    'count': 0,
                    'total': 0
                }
            students_passed[student.id]['records'].append(record)
            students_passed[student.id]['count'] += 1
            students_passed[student.id]['total'] += total_avg
    
    # Calculate overall average for each student
    for student_data in students_passed.values():
        student_data['average'] = student_data['total'] / student_data['count'] if student_data['count'] > 0 else 0
    
    # Sort by average
    sorted_students = sorted(students_passed.values(), key=lambda x: x['average'], reverse=True)
    
    # Create Word document
    doc = Document()
    doc.add_heading('Pass List', 0)
    
    if current_campus:
        doc.add_paragraph(f'Campus: {current_campus.name}')
    elif request.user.is_superuser:
        doc.add_paragraph('All Campuses')
    
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Students Passed: {len(sorted_students)}')
    
    # Add table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Student Name'
    header_cells[1].text = 'Registration Number'
    header_cells[2].text = 'Course'
    header_cells[3].text = 'Average Score'
    
    # Data rows
    for student_data in sorted_students:
        student = student_data['student']
        row_cells = table.add_row().cells
        row_cells[0].text = student.name
        row_cells[1].text = student.registration_number
        row_cells[2].text = student.course.name
        row_cells[3].text = f"{student_data['average']:.2f}"
    
    # Save document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=pass_list.docx'
    doc.save(response)
    return response


def download_records_word(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    if current_campus:
        records = ExamRecord.objects.select_related('student', 'unit', 'unit__course').filter(student__course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            records = ExamRecord.objects.select_related('student', 'unit', 'unit__course').all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    # Create Word document
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Header Image
    try:
        doc.add_picture('C:/Users/PETER/Desktop/exam mangement/report/head.jpg', width=Inches(6.5))
    except FileNotFoundError:
        doc.add_paragraph("Header image not found.")

    # Form details
    p = doc.add_paragraph()
    p.add_run('COURSE: ').bold = True
    p.add_run(f"{request.GET.get('course', '....................................................')}")
    p.add_run('    SEMESTER: ').bold = True
    p.add_run(f"{request.GET.get('term', '................')}")
    p.add_run('    YEAR: ').bold = True
    p.add_run(f"{request.GET.get('year', '................')}")
    
    p = doc.add_paragraph()
    p.add_run('SUBJECT: ').bold = True
    p.add_run(f"{request.GET.get('unit', '..................................................')}")

    # Table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STUDENT NAME'
    hdr_cells[1].text = 'ADM NO'
    hdr_cells[2].text = 'ASSN'
    hdr_cells[3].text = 'CAT 1'
    hdr_cells[4].text = 'CAT 2'
    hdr_cells[5].text = 'END TERM'
    hdr_cells[6].text = 'TOTAL'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table Body (with records)
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = record.student.name.upper()
        row_cells[1].text = record.student.registration_number
        row_cells[2].text = '' # ASSN is not in the model
        row_cells[3].text = str(record.cat1_score)
        row_cells[4].text = str(record.cat2_score)
        row_cells[5].text = str(record.end_term_score)
        row_cells[6].text = str(record.total_average)

    # Add 15 empty rows for manual entry
    for _ in range(15):
        table.add_row()

    # --- Serve the document ---
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    response = HttpResponse(
        f.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = "Marks_Entry_Sheet.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response 

def enter_marks_spreadsheet(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')

    if current_campus:
        courses = Course.objects.filter(school__campus=current_campus)
        units = Unit.objects.filter(course__school__campus=current_campus)
    else:
        courses = Course.objects.all()
        units = Unit.objects.all()

    context = {
        'courses': courses,
        'units': units,
        'current_campus': current_campus,
    }
    return render(request, 'exams/enter_marks_spreadsheet.html', context)


def enter_marks_per_student(request):
    current_campus = get_current_campus(request)
    if not current_campus and not request.user.is_superuser:
        return redirect('exams:campus_select')
    
    # Respect campus selection for both superusers and regular users
    if current_campus:
        students = Student.objects.filter(course__school__campus=current_campus)
    else:
        # If no campus is selected and user is superuser, show all
        if request.user.is_superuser:
            students = Student.objects.all()
        else:
            # Regular users must have a campus selected
            return redirect('exams:campus_select')
    
    years = [2025, 2024, 2023, 2022, 2021]
    terms = ['Term 1', 'Term 2', 'Term 3', 'Term 4']
    
    selected_student = None
    selected_year = None
    selected_term = None
    unit_marks = []
    available_units = []
    message = ''
    
    if request.method == 'POST' and 'select_student' in request.POST:
        student_id = request.POST.get('student_id')
        year = request.POST.get('year')
        term = request.POST.get('term')
        selected_student = get_object_or_404(Student, id=student_id)
        selected_year = year
        selected_term = term
        all_units = list(Unit.objects.filter(course=selected_student.course))
        entered_units = []
        for unit in all_units:
            rec = ExamRecord.objects.filter(student=selected_student, unit=unit, year=year, term=term).first()
            if rec:
                unit_marks.append({
                    'unit': unit,
                    'cat1': rec.cat1_score,
                    'cat2': rec.cat2_score,
                    'endterm': rec.end_term_score,
                    'has_record': True,
                })
                entered_units.append(unit.id)
        available_units = [u for u in all_units if u.id not in entered_units]
    elif request.method == 'POST' and 'save_marks' in request.POST:
        student_id = request.POST.get('student_id')
        year = request.POST.get('year')
        term = request.POST.get('term')
        selected_student = get_object_or_404(Student, id=student_id)
        selected_year = year
        selected_term = term
        school = request.POST.get('school')
        level = request.POST.get('level')
        i = 1
        entered_units = []
        while True:
            unit_id = request.POST.get(f'unit_id_{i}')
            unit_name = request.POST.get(f'unit_name_{i}', '').strip()
            if not unit_id:
                break
            cat1 = request.POST.get(f'cat1_{i}')
            cat2 = request.POST.get(f'cat2_{i}')
            endterm = request.POST.get(f'endterm_{i}')
            cat1_val = safe_decimal(cat1)
            cat2_val = safe_decimal(cat2)
            endterm_val = safe_decimal(endterm)
            if unit_id == 'other' and unit_name:
                unit, _ = Unit.objects.get_or_create(
                    name=unit_name,
                    course=selected_student.course
                )
            else:
                unit = Unit.objects.get(id=unit_id)
            record, _ = ExamRecord.objects.get_or_create(
                student=selected_student,
                unit=unit,
                year=year,
                term=term,
                defaults={
                    'cat1_score': cat1_val,
                    'cat2_score': cat2_val,
                    'end_term_score': endterm_val,
                }
            )
            record.cat1_score = cat1_val
            record.cat2_score = cat2_val
            record.end_term_score = endterm_val
            record.school = school
            record.level = level
            record.save()
            entered_units.append(unit.id)
            i += 1
        message = 'Marks saved successfully!'
        # Reset form for next student
        selected_student = None
        selected_year = None
        selected_term = None
        unit_marks = []
        available_units = []
    context = {
        'students': students,
        'years': years,
        'terms': terms,
        'selected_student': selected_student,
        'selected_year': selected_year,
        'selected_term': selected_term,
        'unit_marks': unit_marks,
        'available_units': available_units,
        'message': message,
        'current_campus': current_campus,
    }
    return render(request, 'exams/enter_marks_per_student.html', context) 

def safe_decimal(val):
    try:
        return float(val)
    except (TypeError, ValueError):
        return 0
